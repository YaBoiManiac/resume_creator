"""Cover letter builder for creating formatted Word documents."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
from typing import Dict


class CoverLetterBuilder:
    """Build professionally formatted cover letters in DOCX format."""
    
    def __init__(self):
        """Initialize the cover letter builder."""
        self.doc = Document()
        self._setup_styles()
        self.output_dir = Path(__file__).parent.parent / "output"
        self.output_dir.mkdir(exist_ok=True)
    
    def _setup_styles(self):
        """Set up document-wide styles and formatting."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
    
    def add_header(self, personal_info: Dict[str, str]):
        """
        Add formatted header with personal information.
        
        Args:
            personal_info: Dictionary containing name, email, phone, etc.
        """
        # Name
        name_paragraph = self.doc.add_paragraph()
        name_run = name_paragraph.add_run(personal_info.get('name', ''))
        name_run.font.size = Pt(14)
        name_run.font.bold = True
        
        # Contact information
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        
        if contact_parts:
            contact_paragraph = self.doc.add_paragraph()
            contact_run = contact_paragraph.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(10)
        
        # Date
        date_paragraph = self.doc.add_paragraph()
        date_paragraph.paragraph_format.space_before = Pt(12)
        date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(11)
        
        # Add spacing after header
        self.doc.add_paragraph()
    
    def add_body(self, cover_letter_text: str):
        """
        Add the cover letter body text.
        
        Args:
            cover_letter_text: The generated cover letter content
        """
        # Split into paragraphs
        paragraphs = cover_letter_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                paragraph = self.doc.add_paragraph(para_text.strip())
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.line_spacing = 1.15
                
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    def add_closing(self, name: str):
        """
        Add professional closing.
        
        Args:
            name: Candidate's name
        """
        # Spacing before closing
        self.doc.add_paragraph()
        
        # Closing
        closing_paragraph = self.doc.add_paragraph("Best regards,")
        closing_paragraph.paragraph_format.space_after = Pt(24)
        
        for run in closing_paragraph.runs:
            run.font.size = Pt(11)
        
        # Name
        name_paragraph = self.doc.add_paragraph(name)
        for run in name_paragraph.runs:
            run.font.size = Pt(11)
    
    def save(self, filename: str = None) -> Path:
        """
        Save the cover letter to a DOCX file.
        
        Args:
            filename: Optional custom filename (without extension)
            
        Returns:
            Path to the saved file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"cover_letter_{timestamp}"
        
        # Add _cover_letter suffix if not present
        if not 'cover' in filename.lower():
            filename = filename + "_cover_letter"
        
        # Ensure .docx extension
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        filepath = self.output_dir / filename
        self.doc.save(str(filepath))
        
        return filepath
    
    def build_complete_cover_letter(
        self,
        personal_info: Dict[str, str],
        cover_letter_text: str,
        filename: str = None
    ) -> Path:
        """
        Build a complete cover letter.
        
        Args:
            personal_info: Personal information dictionary
            cover_letter_text: The generated cover letter content
            filename: Optional custom filename
            
        Returns:
            Path to the saved cover letter file
        """
        self.add_header(personal_info)
        self.add_body(cover_letter_text)
        self.add_closing(personal_info.get('name', ''))
        
        return self.save(filename)
