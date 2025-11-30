"""Resume builder for creating professionally formatted Word documents."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any


class ResumeBuilder:
    """Build professionally formatted resumes in DOCX format."""
    
    def __init__(self):
        """Initialize the resume builder."""
        self.doc = Document()
        self._setup_styles()
        self.output_dir = Path(__file__).parent.parent / "output"
        self.output_dir.mkdir(exist_ok=True)
    
    def _setup_styles(self):
        """Set up document-wide styles and formatting."""
        # Set narrow margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
    
    def add_header(self, personal_info: Dict[str, str]):
        """
        Add formatted header with personal information.
        
        Args:
            personal_info: Dictionary containing name, email, phone, location, etc.
        """
        # Name - large and bold
        name_paragraph = self.doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(personal_info.get('name', ''))
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Contact information - centered
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        
        if contact_parts:
            contact_paragraph = self.doc.add_paragraph()
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_paragraph.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(10)
        
        # LinkedIn and Portfolio - centered
        links = []
        if personal_info.get('linkedin'):
            links.append(personal_info['linkedin'])
        if personal_info.get('portfolio'):
            links.append(personal_info['portfolio'])
        
        if links:
            links_paragraph = self.doc.add_paragraph()
            links_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_run = links_paragraph.add_run(' | '.join(links))
            links_run.font.size = Pt(9)
            links_run.font.color.rgb = RGBColor(0, 0, 255)
        
        # Add spacing after header
        self.doc.add_paragraph()
    
    def add_section_heading(self, title: str):
        """
        Add a section heading (e.g., "Professional Summary", "Experience").
        
        Args:
            title: The section title
        """
        heading = self.doc.add_paragraph()
        heading_run = heading.add_run(title.upper())
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add a line under the heading
        heading.paragraph_format.space_after = Pt(6)
    
    def add_about_me(self, about_me_text: str):
        """
        Add the "About Me" or "Professional Summary" section.
        
        Args:
            about_me_text: The generated about me text
        """
        self.add_section_heading("Professional Summary")
        
        paragraph = self.doc.add_paragraph(about_me_text)
        paragraph.paragraph_format.space_after = Pt(12)
        
        for run in paragraph.runs:
            run.font.size = Pt(10)
    
    def add_job_experience(self, jobs: List[Dict[str, Any]], enhanced_duties: Dict[int, List[str]] = None):
        """
        Add job experience section with formatted entries.
        
        Args:
            jobs: List of job experience dictionaries
            enhanced_duties: Optional dict mapping job indices to enhanced duty lists
        """
        if not jobs:
            return
        
        self.add_section_heading("Professional Experience")
        
        for idx, job in enumerate(jobs):
            # Job title and company - bold
            title_paragraph = self.doc.add_paragraph()
            title_run = title_paragraph.add_run(
                f"{job.get('title', 'N/A')} | {job.get('company', 'N/A')}"
            )
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            
            # Dates and location
            date_paragraph = self.doc.add_paragraph()
            date_text = f"{job.get('start_date', 'N/A')} - {job.get('end_date', 'N/A')}"
            date_run = date_paragraph.add_run(date_text)
            date_run.font.size = Pt(10)
            date_run.font.italic = True
            
            # Use enhanced duties if provided, otherwise use original
            duties = enhanced_duties.get(idx) if enhanced_duties else job.get('duties', [])
            
            for duty in duties:
                duty_paragraph = self.doc.add_paragraph(duty, style='List Bullet')
                duty_paragraph.paragraph_format.left_indent = Inches(0.25)
                duty_paragraph.paragraph_format.space_after = Pt(2)
                
                for run in duty_paragraph.runs:
                    run.font.size = Pt(10)
            
            # Add space after each job
            self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    def add_education(self, education: List[Dict[str, str]]):
        """
        Add education section.
        
        Args:
            education: List of education entries
        """
        if not education:
            return
        
        self.add_section_heading("Education")
        
        for edu in education:
            edu_paragraph = self.doc.add_paragraph()
            
            # Degree and institution
            degree_run = edu_paragraph.add_run(
                f"{edu.get('degree', 'N/A')} - {edu.get('institution', 'N/A')}"
            )
            degree_run.font.size = Pt(10)
            degree_run.font.bold = True
            
            # Year
            if edu.get('year'):
                edu_paragraph.add_run(f" ({edu['year']})")
            
            edu_paragraph.paragraph_format.space_after = Pt(4)
        
        self.doc.add_paragraph()
    
    def add_skills(self, skills: List[str]):
        """
        Add skills section.
        
        Args:
            skills: List of skills
        """
        if not skills:
            return
        
        self.add_section_heading("Skills")
        
        skills_text = " • ".join(skills)
        skills_paragraph = self.doc.add_paragraph(skills_text)
        skills_paragraph.paragraph_format.space_after = Pt(12)
        
        for run in skills_paragraph.runs:
            run.font.size = Pt(10)
    
    def add_certifications(self, certifications: List[str]):
        """
        Add certifications section.
        
        Args:
            certifications: List of certifications
        """
        if not certifications:
            return
        
        self.add_section_heading("Certifications")
        
        for cert in certifications:
            cert_paragraph = self.doc.add_paragraph(cert, style='List Bullet')
            cert_paragraph.paragraph_format.left_indent = Inches(0.25)
            cert_paragraph.paragraph_format.space_after = Pt(2)
            
            for run in cert_paragraph.runs:
                run.font.size = Pt(10)
    
    def save(self, filename: str = None) -> Path:
        """
        Save the resume to a DOCX file.
        
        Args:
            filename: Optional custom filename (without extension)
            
        Returns:
            Path to the saved file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}"
        
        # Ensure .docx extension
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        filepath = self.output_dir / filename
        self.doc.save(str(filepath))
        
        return filepath
    
    def build_complete_resume(
        self,
        personal_info: Dict[str, str],
        about_me: str,
        jobs: List[Dict[str, Any]],
        enhanced_duties: Dict[int, List[str]] = None,
        education: List[Dict[str, str]] = None,
        skills: List[str] = None,
        certifications: List[str] = None,
        filename: str = None
    ) -> Path:
        """
        Build a complete resume with all sections.
        
        Args:
            personal_info: Personal information dictionary
            about_me: Professional summary text
            jobs: List of job experiences
            enhanced_duties: Optional enhanced duties for jobs
            education: Optional education list
            skills: Optional skills list
            certifications: Optional certifications list
            filename: Optional custom filename
            
        Returns:
            Path to the saved resume file
        """
        self.add_header(personal_info)
        self.add_about_me(about_me)
        self.add_job_experience(jobs, enhanced_duties)
        
        if education:
            self.add_education(education)
        
        if skills:
            self.add_skills(skills)
        
        if certifications:
            self.add_certifications(certifications)
        
        return self.save(filename)
